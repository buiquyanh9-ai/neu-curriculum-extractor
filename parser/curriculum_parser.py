"""
parser/curriculum_parser.py — Rule-based CTĐT extractor.
Hỗ trợ cả 3 dạng format:
  - Chính quy tiếng Việt (Bảo hiểm, KDQT, Thống kê...)
  - AEP/CLC tiếng Anh (FE, EPMP, EDA, IHME...)
  - Form mới 2025 (Phân tích dữ liệu kinh tế...)
"""
import re
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.table import Table

from .models import (
    CurriculumRecord,
    ProgramObjective, ProgramLearningOutcome, POPLOMap,
    ProgramSection, ProgramComponent,
    ProgramCourse, CourseSemesterAllocation,
    CareerPath, GraduationRequirement, ProgramReferenceProgram,
)

# ── Helpers ───────────────────────────────────────────────────────────────────

def _c(t: str) -> str:
    """Clean whitespace."""
    return re.sub(r"\s+", " ", (t or "").replace("\u200b", "").replace("\xa0", " ")).strip()

def _lines(text: str) -> list:
    return [_c(l) for l in text.splitlines() if _c(l)]

def _all_para_text(doc: Document) -> list:
    return [_c(p.text) for p in doc.paragraphs if _c(p.text)]

def _split_vi_en(text: str) -> tuple:
    """
    Tách tên VI / EN: 'Triết học Mác - Lê Nin  Marxist Philosophy'
    hoặc 'Triết học / Philosophy', 'Triết học\nPhilosophy'.
    Trả về (vi, en).
    """
    # Dấu gạch chéo với khoảng trắng: "Tên / Name"
    m = re.split(r"\s{2,}|\n|(?<=\w)\s*/\s*(?=[A-Z])", text, maxsplit=1)
    if len(m) == 2:
        vi = _c(m[0]); en = _c(m[1])
        # Kiểm tra phần sau có vẻ là tiếng Anh không
        if re.search(r"[a-z]{2}", en):
            return vi, en
    return text, None

def _parse_credit(text: str) -> Optional[float]:
    text = _c(text)
    if text in ("-", "", "—", "–", "TC", "None"): return None
    m = re.search(r"(\d+(?:[.,]\d+)?)", text)
    if m:
        try:
            v = float(m.group(1).replace(",", "."))
            if 0 < v <= 300: return v
        except: pass
    return None

def _roman_semester(s: str) -> list:
    """
    Chuyển 'I-II', 'III=>VI', '1', '2,3', 'I,II,III' → list[int].
    """
    ROMAN = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6,
             "VII": 7, "VIII": 8, "IX": 9, "X": 10}
    s = _c(s).upper()
    if not s or s in ("-", "—"): return []
    result = set()
    # Thay thế mũi tên
    s = re.sub(r"=>|→|ĐẾN|TO", "-", s, flags=re.I)
    # Xử lý từng phần cách nhau bởi dấu phẩy
    for part in re.split(r"[,;]", s):
        part = part.strip()
        # Dạng X-Y hoặc X=> Y (range)
        range_m = re.match(r"^(I{1,3}V?|V?I{0,3}|\d+)\s*[-–]\s*(I{1,3}V?|V?I{0,3}|\d+)$", part)
        if range_m:
            a_str, b_str = range_m.group(1), range_m.group(2)
            a = ROMAN.get(a_str, None) or (int(a_str) if a_str.isdigit() else None)
            b = ROMAN.get(b_str, None) or (int(b_str) if b_str.isdigit() else None)
            if a and b:
                for i in range(min(a,b), max(a,b)+1):
                    if 1 <= i <= 12: result.add(i)
            continue
        # Đơn: roman hoặc số
        v = ROMAN.get(part, None)
        if v is None and part.isdigit():
            v = int(part)
        if v and 1 <= v <= 12:
            result.add(v)
    return sorted(result)


# ── Thông tin chung ───────────────────────────────────────────────────────────

_ISSUED_DECISION_RE = re.compile(
    r"(?:ban\s*hành\s*theo\s*quyết\s*định\s*số|theo\s*quyết\s*định\s*số|quyết\s*định\s*số)\s*(.+?)(?:\s*của|\s*ngày|\s*$)",
    re.I | re.S,
)

def _parse_general_info(doc: Document) -> dict:
    info = dict(
        program_name_vi=None, program_name_en=None,
        degree_name_vi=None, degree_name_en=None,
        level_name_vi=None, level_name_en=None,
        major_code=None, major_name_vi=None, major_name_en=None,
        training_type_vi=None, training_type_en=None,
        language_vi=None, language_en=None,
        duration_text_vi=None, duration_text_en=None,
        total_credits=None,
        applied_from_admission_year=None,
        adjustment_time_text=None,
        program_opening_decision_text=None,
        accreditation_text=None,
        degree_granting_unit_text=None,
        academic_management_unit_text=None,
        issued_decision_text=None,
        version_label=None,
    )

    # Tìm "issued decision" từ paragraphs đầu
    for p in doc.paragraphs[:30]:
        t = _c(p.text)
        if not t: continue
        m = _ISSUED_DECISION_RE.search(t)
        if m:
            info["issued_decision_text"] = _c(m.group(0))
            # Tìm năm để làm version label
            yr_m = re.search(r"\b(20\d{2})\b", t)
            if yr_m: info["version_label"] = yr_m.group(1)
            break

    # Tìm total_credits từ bảng "KHỐI LƯỢNG KIẾN THỨC TOÀN KHOÁ | | 123 tín chỉ"
    for tbl in doc.tables[:8]:
        for row in tbl.rows:
            cells = [_c(c.text) for c in row.cells]
            row_text = " ".join(cells)
            if re.search(r"khối\s*lượng\s*kiến\s*thức|total\s*credits?|tổng\s*số\s*tín\s*chỉ", row_text, re.I):
                for cell in cells:
                    m2 = re.search(r"(\d+)\s*tín\s*chỉ", cell, re.I)
                    if not m2:
                        m2 = re.search(r"^(\d{2,3})$", cell)
                    if m2:
                        v = float(m2.group(1))
                        if 60 <= v <= 300:
                            info["total_credits"] = v
                            break
                if info["total_credits"]:
                    break
        if info["total_credits"]:
            break

    # Tìm table thông tin chung (thường là table 0-2)
    for tbl in doc.tables[:6]:
        rows = tbl.rows
        if len(rows) < 3: continue
        # Kiểm tra xem có phải bảng thông tin chung không
        all_cells = " ".join(_c(c.text).lower() for r in rows for c in r.cells)
        if not re.search(r"chương\s*trình|programme|trình\s*độ|level", all_cells, re.I):
            continue

        for row in rows:
            cells = [_c(c.text) for c in row.cells]
            # Merge duplicated cells (docx merged cells show same text)
            unique = list(dict.fromkeys(cells))
            if len(unique) < 2: continue
            label, value = unique[0], " ".join(unique[1:])

            # Normalize label — strip "(PROGRAMME)", "(CODE)" etc. bilingual suffixes
            lbl = re.sub(r"\s*\([^)]*\)\s*$", "", label.lower()).strip().rstrip(":")

            # Program / Degree name
            # Exclude "kiểm định chương trình", "điều chỉnh chương trình" rows
            # Only set once (don't overwrite a good value)
            if (re.search(r"chương\s*trình|programme", lbl)
                    and not re.search(r"kiểm\s*định|điều\s*chỉnh|accredit|adjust|tham\s*khảo", lbl)
                    and not info["program_name_vi"]):
                vi, en = _split_vi_en(value)
                # Value must look like a real programme name (not a year or short text)
                if vi and len(vi) > 5 and not re.match(r"^(năm|tháng|\d)", vi, re.I):
                    info["program_name_vi"] = vi
                    if en: info["program_name_en"] = en
                    dm = re.match(r"(cử\s*nhân|thạc\s*sĩ|tiến\s*sĩ)\s+(.+)", vi, re.I)
                    if dm:
                        info["degree_name_vi"] = _c(dm.group(1)).title() + " " + _c(dm.group(2))


            # Level
            elif re.search(r"trình\s*độ|level\s*of\s*edu", lbl):
                vi, en = _split_vi_en(value)
                info["level_name_vi"] = vi
                if en: info["level_name_en"] = en

            # Major (Ngành đào tạo) — exclude mã ngành
            elif re.search(r"ngành\s*đào\s*tạo|^major$", lbl) and "mã" not in lbl:
                vi, en = _split_vi_en(value)
                info["major_name_vi"] = vi
                if en: info["major_name_en"] = en

            # Major code — "mã ngành đào tạo", "code"
            elif re.search(r"mã\s*ngành|mã\s*đào\s*tạo|^code$", lbl):
                codes = re.findall(r"\b(\d{7})\b", value)
                if not codes:
                    codes = re.findall(r"\b([A-Z]{2,6}\d{1,3})\b", value, re.I)
                if codes:
                    info["major_code"] = codes[0]
                elif value.strip():
                    info["major_code"] = _c(value)

            # Training type
            elif re.search(r"hình\s*thức|type\s*of\s*edu", lbl):
                vi, en = _split_vi_en(value)
                info["training_type_vi"] = vi
                if en: info["training_type_en"] = en

            # Language
            elif re.search(r"ngôn\s*ngữ|^language$", lbl):
                vi, en = _split_vi_en(value)
                info["language_vi"] = vi
                if en: info["language_en"] = en

            # Duration
            elif re.search(r"thời\s*gian|duration", lbl):
                vi, en = _split_vi_en(value)
                info["duration_text_vi"] = vi
                if en: info["duration_text_en"] = en

            # Total credits
            elif re.search(r"số\s*tín\s*chỉ|total\s*credits?", lbl):
                v = _parse_credit(value)
                if v and not info["total_credits"]: info["total_credits"] = v

            # Opening decision
            elif re.search(r"quyết\s*định\s*mở|opening\s*decision|cấp\s*ban\s*hành", lbl):
                info["program_opening_decision_text"] = _c(value)

            # Accreditation
            elif re.search(r"kiểm\s*định", lbl):
                info["accreditation_text"] = _c(value)

            # Adjustment time
            elif re.search(r"điều\s*chỉnh", lbl):
                info["adjustment_time_text"] = _c(value)

            # Applied from year
            elif re.search(r"áp\s*dụng|tuyển\s*sinh", lbl):
                yr_m = re.search(r"\b(20\d{2})\b", value)
                if yr_m: info["applied_from_admission_year"] = int(yr_m.group(1))
                if not info["version_label"]: info["version_label"] = _c(value)

            # Degree granting unit
            elif re.search(r"đơn\s*vị\s*cấp\s*bằng|granting", lbl):
                info["degree_granting_unit_text"] = _c(value)

            # Academic management unit
            elif re.search(r"đơn\s*vị\s*quản\s*lý|academic\s*management", lbl):
                info["academic_management_unit_text"] = _c(value)

        # Nếu đã lấy được program_name, coi như xong bảng này
        if info["program_name_vi"]:
            break

    # Fallback: lấy chương trình từ tiêu đề đậm trong paragraphs
    if not info["program_name_vi"]:
        for p in doc.paragraphs[:50]:
            t = _c(p.text)
            if not t: continue
            if re.search(r"^(cử\s*nhân|thạc\s*sĩ|bachelor)\s", t, re.I) and len(t) < 120:
                vi, en = _split_vi_en(t)
                info["program_name_vi"] = vi
                if en: info["program_name_en"] = en
                break

    if not info["version_label"]:
        info["version_label"] = "unknown"

    return info


# ── Philosophy ────────────────────────────────────────────────────────────────

_PHIL_HDR = re.compile(r"triết\s*lý\s*giáo\s*dục|education\s*philosophy|philosophy", re.I)
_GOAL_HDR = re.compile(
    r"mục\s*tiêu\s*(?:đào\s*tạo|chương\s*trình)"  # "Mục tiêu đào tạo"
    r"|objectives?"                                    # "Objectives"
    r"|goals?\s*of\s*train"                         # "Goals of training"
    r"|^\d+\.\s*mục\s*tiêu\s*$"                 # "1. MỤC TIÊU" (numbered)
    r"|^mục\s*tiêu\s*$",                            # "MỤC TIÊU" plain (FE format)
    re.I,
)
_PLO_HDR  = re.compile(r"chuẩn\s*đầu\s*ra|learning\s*outcomes?|graduate\s*attributes?|plo", re.I)
_CAREER_HDR = re.compile(
    r"(?:cơ\s*hội\s*(?:việc\s*làm|nghề\s*nghiệp|thăng\s*tiến)"
    r"|định\s*hướng\s*(?:nơi\s*làm\s*việc|sau\s*tốt\s*nghiệp)"
    r"|nơi\s*làm\s*việc\s*sau\s*tốt\s*nghiệp"
    r"|vị\s*trí\s*(?:và\s*nơi|làm\s*việc|công\s*việc)"
    r"|việc\s*làm\s*sau\s*tốt\s*nghiệp"
    r"|mục\s*tiêu\s*sau\s*khi\s*tốt\s*nghiệp"
    r"|career\s*(?:opportunity|path|prospect)"
    r"|employment\s*(?:opportunity|prospect)"
    r"|after\s*graduation)",
    re.I,
)
_GRAD_REQ_HDR = re.compile(r"điều\s*kiện\s*tốt\s*nghiệp|graduation\s*req|yêu\s*cầu\s*tốt\s*nghiệp", re.I)
_STRUCT_HDR   = re.compile(r"cấu\s*trúc\s*chương\s*trình|khối\s*lượng\s*kiến\s*thức|structure", re.I)
_REF_PROG_HDR = re.compile(r"chương\s*trình\s*tham\s*khảo|reference\s*program", re.I)
_STAFF_HDR    = re.compile(r"tiêu\s*chuẩn\s*giảng\s*viên|staff\s*standard", re.I)
_FACILITY_HDR = re.compile(r"cơ\s*sở\s*vật\s*chất|facility|facilities", re.I)
_GUIDE_HDR    = re.compile(r"hướng\s*dẫn\s*thực\s*hiện|implementation\s*guide", re.I)

def _parse_philosophy(doc: Document) -> tuple:
    """Trả về (vi_text, en_text) của triết lý giáo dục."""
    collecting, chunks_vi, chunks_en = False, [], []
    for p in doc.paragraphs:
        t = _c(p.text)
        if not t: continue
        if _PHIL_HDR.search(t) and len(t) < 80:
            collecting = True; continue
        if collecting:
            if _GOAL_HDR.search(t) and len(t) < 80: break
            if re.match(r"^(\d+\.|\*|–)\s", t) and len(t) < 60: break
            if re.search(r"[a-z]{4}", t) and not re.search(r"[àáâãèéêìíòóôõùúăắặằẵặêếệềểẽôốộồổỗơớợờởỡưứựừửữ]", t):
                chunks_en.append(t)
            else:
                chunks_vi.append(t)
    return " ".join(chunks_vi) or None, " ".join(chunks_en) or None


# ── Objectives (PO) ───────────────────────────────────────────────────────────

def _parse_objectives(doc: Document) -> list:
    objectives = []

    # 1) Từ bảng PO
    for tbl in doc.tables:
        if not tbl.rows: continue
        hdr_text = " ".join(_c(c.text).lower() for c in tbl.rows[0].cells)
        if not re.search(r"mã|code|^po\d?$|mục\s*tiêu|objective", hdr_text): continue
        if not re.search(r"mô\s*tả|description|objective|nội\s*dung", hdr_text): continue
        # Phải là bảng PO chứ không phải PLO
        # Bỏ qua bảng PLO, bảng khóa trình, bảng danh sách HP
        if re.search(r"plo|chuẩn\s*đầu\s*ra\b|mã\s*hp|học\s*phần", hdr_text, re.I): continue

        code_col = desc_col = -1
        for i, h in enumerate(_c(c.text).lower() for c in tbl.rows[0].cells):
            if re.search(r"^mã$|^code$|^stt$|^tt$|^po", h) and code_col < 0: code_col = i
            if re.search(r"mô\s*tả|description|nội\s*dung|objective|mục\s*tiêu", h) and i != code_col and desc_col < 0:
                desc_col = i

        if code_col < 0 and desc_col >= 0: code_col = 0
        if code_col < 0 or desc_col < 0: continue

        order = 1
        for row in tbl.rows[1:]:
            cells = [_c(c.text) for c in row.cells]
            if not cells: continue
            if code_col >= len(cells) or desc_col >= len(cells): continue
            code = cells[code_col]
            desc = cells[desc_col]
            if not desc or len(desc) < 5: continue

            # Normalize code: PO1, G1, etc.
            po_m = re.match(r"^(PO\s*\d+|G\s*\d+|\d+)$", code, re.I)
            if po_m:
                norm_code = re.sub(r"\s+", "", po_m.group(1).upper())
                if norm_code.isdigit(): norm_code = f"PO{norm_code}"
            elif re.match(r"mục\s*tiêu\s*chung|general", code, re.I):
                norm_code = "PO_GENERAL"
            else:
                continue

            objectives.append(ProgramObjective(
                objective_code=norm_code,
                objective_type="specific" if re.match(r"PO\d+|G\d+", norm_code) else "general",
                description_vi=desc,
                display_order=order,
            ))
            order += 1
        if objectives: break

    # 2) Fallback: từ paragraphs dưới heading "Mục tiêu đào tạo"
    if not objectives:
        objectives = _parse_objectives_from_paragraphs(doc)

    return objectives


def _parse_objectives_from_paragraphs(doc: Document) -> list:
    objectives = []
    collecting = False
    order = 1
    general_chunks = []

    for p in doc.paragraphs:
        t = _c(p.text)
        if not t: continue

        if _GOAL_HDR.search(t) and len(t) < 100:
            collecting = True; continue

        if collecting:
            # Dừng khi gặp section khác
            if (_PLO_HDR.search(t) or _STRUCT_HDR.search(t)) and len(t) < 80:
                break
            # Dừng ở heading số cấp 1: "2. CHUẨN ĐẦU RA", "7. CẤU TRÚC..."
            if re.match(r"^\d+\.\s{1,3}\S", t) and len(t) < 80 and not _GOAL_HDR.search(t):
                break
            # Career info → skip (đừng thu vào objectives)
            if _CAREER_HDR.search(t): continue
            # Numbered PO: "PO1:", "1.", "- Objective 1:"
            po_m = re.match(r"^(?:(PO\s*\d+)[:\.]?\s*|(\d+)\.\s+(?=[A-ZÁÀẢÃẠĂẮẶẰẴẶÊẾỆỀỂẼÔỐỘỒỔỖƠỚỢỜỞỠƯỨỰỪỬỮ]))", t, re.I)
            if po_m:
                code = re.sub(r"\s+", "", (po_m.group(1) or f"PO{po_m.group(2)}").upper())
                desc = t[po_m.end():].strip()
                if desc:
                    objectives.append(ProgramObjective(
                        objective_code=code,
                        objective_type="specific",
                        description_vi=desc,
                        display_order=order,
                    ))
                    order += 1
            elif len(t) > 30:
                general_chunks.append(t)

    # Thêm objective chung nếu có text
    if general_chunks and not objectives:
        objectives.insert(0, ProgramObjective(
            objective_code="PO_GENERAL",
            objective_type="general",
            description_vi=" ".join(general_chunks[:5]),
            display_order=0,
        ))

    return objectives


# ── PLO ───────────────────────────────────────────────────────────────────────

_PLO_CODE_RE = re.compile(
    r"^(PLO|CDR|CĐR|ELO|LO)\s*\.?\s*(\d+)(?:[.\-](\d+))?$",
    re.I,
)
_PLO_NUM_RE = re.compile(r"^(\d+)(?:[.\-](\d+))?$")

_PLO_GROUP_MAP = {
    "kiến thức": ("1", "Kiến thức", "Knowledge"),
    "kỹ năng": ("2", "Kỹ năng", "Skills"),
    "tự chủ": ("3", "Tự chủ và trách nhiệm", "Autonomy and Responsibility"),
    "trách nhiệm": ("3", "Tự chủ và trách nhiệm", "Autonomy and Responsibility"),
    "knowledge": ("1", "Kiến thức", "Knowledge"),
    "skill": ("2", "Kỹ năng", "Skills"),
    "autonomy": ("3", "Tự chủ và trách nhiệm", "Autonomy and Responsibility"),
}

def _guess_plo_group(code: str, group_name: str) -> tuple:
    """Đoán nhóm PLO từ code hoặc tên nhóm. Trả về (group_code, name_vi, name_en)."""
    text = (code + " " + group_name).lower()
    for kw, vals in _PLO_GROUP_MAP.items():
        if kw in text:
            return vals
    # Dựa vào số đầu của code
    first_digit_m = re.search(r"(\d+)", code)
    if first_digit_m:
        d = first_digit_m.group(1)
        defaults = {
            "1": ("1", "Kiến thức", "Knowledge"),
            "2": ("2", "Kỹ năng", "Skills"),
            "3": ("3", "Tự chủ và trách nhiệm", "Autonomy"),
        }
        if d in defaults: return defaults[d]
    return ("0", group_name or "Other", None)


def _parse_plos(doc: Document) -> list:
    plos = []

    for tbl in doc.tables:
        if not tbl.rows or len(tbl.rows) < 3: continue
        hdr_text = " ".join(_c(c.text).lower() for c in tbl.rows[0].cells)

        # Phải là bảng PLO
        if not re.search(r"plo|chuẩn\s*đầu\s*ra|learning\s*outcome|elo", hdr_text, re.I):
            continue
        # Không phải bảng PO-PLO matrix (có nhiều cột PLO)
        plo_cols = len(re.findall(r"plo\s*\.?\s*\d+", hdr_text, re.I))
        if plo_cols > 3: continue

        code_col = desc_col = group_col = -1
        headers = [_c(c.text).lower() for c in tbl.rows[0].cells]
        for i, h in enumerate(headers):
            if re.search(r"^mã\s*plo$|^mã$|^code$|^stt$|^tt$|^plo\d*$|^elo$", h) and code_col < 0:
                code_col = i
            if re.search(r"mô\s*tả|description|nội\s*dung", h) and i != code_col and desc_col < 0:
                desc_col = i
            if re.search(r"nhóm|group", h) and group_col < 0:
                group_col = i

        # Dạng AEP: "Chuẩn đầu ra | Mô tả | Mã PLO"
        # Col 0 = group label (repeats per group), Col 1 = description, Col 2 = PLO code
        if code_col < 0 and desc_col < 0 and len(headers) == 3:
            if re.search(r"plo|mã|code", headers[-1]):
                code_col = len(headers) - 1  # last col = PLO code
                desc_col = 1                   # middle col = description
                group_col = 0                  # first col = group heading
        elif code_col < 0 and desc_col >= 0:
            # code might be last col
            for i in range(len(headers)-1, -1, -1):
                if re.search(r"plo|mã|code|^\d+$", headers[i]):
                    code_col = i; break
            if code_col < 0:
                code_col = 0

        if code_col < 0: code_col = 0
        if desc_col < 0: desc_col = min(1, len(headers)-1)

        current_group_vi = ""
        current_group_en = None
        current_group_code = "0"
        order = 1

        for row in tbl.rows[1:]:
            cells = [_c(c.text) for c in row.cells]
            if not cells: continue

            # Detect group heading row: merged row with group name
            unique_vals = list(dict.fromkeys(c for c in cells if c))
            if len(unique_vals) == 1 and len(unique_vals[0]) < 80:
                # Likely a group heading
                gv = unique_vals[0]
                gc, gn_vi, gn_en = _guess_plo_group("", gv)
                current_group_vi = gn_vi or gv
                current_group_en = gn_en
                current_group_code = gc
                continue

            code_raw = cells[code_col] if code_col < len(cells) else ""
            desc_raw = cells[desc_col] if desc_col < len(cells) else ""
            # group_col: detect group from dedicated column OR from code_col text
            # group_col: dùng explicit group_col, nếu không thì thử cells[0]
            if group_col >= 0 and group_col < len(cells):
                group_raw = cells[group_col]
            elif code_col != 0 and desc_col != 0 and len(cells) > 0:
                group_raw = cells[0]  # col0 không phải code/desc → thử làm group
            else:
                group_raw = ""

            # Detect group heading:
            # a) group_col explicitly (AEP format: col0 = "Chuẩn đầu ra về kiến thức")
            if group_raw and re.search(r"kiến\s*thức|kỹ\s*năng|tự\s*chủ|ngoại\s*ngữ|tin\s*học|knowledge|skill|autonomy", group_raw, re.I) and not re.search(r"PLO|\d+\.\d+", group_raw, re.I):
                gc, gn_vi, gn_en = _guess_plo_group("", group_raw)
                current_group_vi = gn_vi or group_raw
                current_group_en = gn_en
                current_group_code = gc
            # b) group in code_col (standard format)
            elif re.search(r"kiến\s*thức|kỹ\s*năng|tự\s*chủ|knowledge|skill|autonomy", code_raw, re.I) and not re.search(r"PLO|CLO|\d+\.\d+", code_raw, re.I):
                gc, gn_vi, gn_en = _guess_plo_group("", code_raw)
                current_group_vi = gn_vi or code_raw
                current_group_en = gn_en
                current_group_code = gc
                continue

            if not code_raw or len(desc_raw) < 5: continue

            # Parse PLO code: PLO1.1, PLO 1, PLO.1, 1.1, 1
            plo_code = None
            pm = _PLO_CODE_RE.match(code_raw)
            if pm:
                major = pm.group(2)
                minor = pm.group(3)
                plo_code = f"PLO{major}.{minor}" if minor else f"PLO{major}"
                # Só infer group nếu chưa có group thực sự từ table
                if not current_group_vi:
                    gc, gn_vi, gn_en = _guess_plo_group(plo_code, "")
                    current_group_code = gc
                    current_group_vi = gn_vi or ""
                    current_group_en = gn_en
            else:
                nm = _PLO_NUM_RE.match(code_raw)
                if nm:
                    major = nm.group(1)
                    minor = nm.group(2)
                    plo_code = f"PLO{major}.{minor}" if minor else f"PLO{major}"
                else:
                    # Check if desc_raw contains a PLO code (inverted cols)
                    inv_m = _PLO_CODE_RE.match(desc_raw)
                    if inv_m:
                        major = inv_m.group(2); minor = inv_m.group(3)
                        plo_code = f"PLO{major}.{minor}" if minor else f"PLO{major}"
                        desc_raw, code_raw = code_raw, desc_raw

            if not plo_code: continue

            plos.append(ProgramLearningOutcome(
                outcome_code=plo_code,
                outcome_group_code=current_group_code,
                outcome_group_name_vi=current_group_vi,
                outcome_group_name_en=current_group_en,
                description_vi=desc_raw,
                display_order=order,
            ))
            order += 1

        if plos: break

    return plos


# ── PO-PLO Matrix ─────────────────────────────────────────────────────────────

def _parse_po_plo_map(doc: Document, objectives: list, plos: list) -> list:
    po_codes  = [o.objective_code for o in objectives]
    plo_codes = [p.outcome_code for p in plos]
    maps = []

    for tbl in doc.tables:
        if not tbl.rows or len(tbl.rows) < 3: continue
        hdr_text = " ".join(_c(c.text).lower() for c in tbl.rows[0].cells)
        # Ma trận PO-PLO: có nhiều cột PLO trong header
        plo_count = len(re.findall(r"plo\s*\.?\s*\d+|\bcđr\s*\d+", hdr_text, re.I))
        if plo_count < 2: continue

        # Parse header row: [PLO] | PLO1 | PLO2 | ...
        hdr_cells = [_c(c.text) for c in tbl.rows[0].cells]
        # Tìm secondary header nếu có
        sec_hdr = [_c(c.text) for c in tbl.rows[1].cells] if len(tbl.rows) > 1 else []
        col_plo_map = {}
        for i, h in enumerate(hdr_cells):
            pm = _PLO_CODE_RE.match(h)
            if pm:
                major = pm.group(2); minor = pm.group(3)
                col_plo_map[i] = f"PLO{major}.{minor}" if minor else f"PLO{major}"
            elif sec_hdr and i < len(sec_hdr):
                pm2 = _PLO_CODE_RE.match(sec_hdr[i])
                if pm2:
                    major = pm2.group(2); minor = pm2.group(3)
                    col_plo_map[i] = f"PLO{major}.{minor}" if minor else f"PLO{major}"

        if not col_plo_map: continue

        # Parse data rows
        start_row = 2 if sec_hdr and any(_PLO_CODE_RE.match(h) for h in sec_hdr) else 1
        for row in tbl.rows[start_row:]:
            cells = [_c(c.text) for c in row.cells]
            if not cells: continue
            row_label = cells[0]
            po_m = re.match(r"^(PLO|PO|CĐR|CDR)\s*\.?\s*(\d+)(?:[.\-](\d+))?$", row_label, re.I)
            if not po_m: continue
            plo_code = f"PLO{po_m.group(2)}.{po_m.group(3)}" if po_m.group(3) else f"PLO{po_m.group(2)}"

            for col_i, po_code in col_plo_map.items():
                if col_i >= len(cells): continue
                val = cells[col_i]
                if val and re.search(r"x|✓|●|■|\d", val, re.I):
                    # Row is PLO, col is PO (or vice versa)
                    maps.append(POPLOMap(
                        po_code=po_code,
                        plo_code=plo_code,
                        mapping_value=val,
                    ))

        if maps: break

    return maps


# ── Sections (narrative) ──────────────────────────────────────────────────────

_SECTION_PATTERNS = [
    (_CAREER_HDR,    "career_opportunity",         "Cơ hội việc làm", "Career Opportunities"),
    (_GRAD_REQ_HDR,  "graduation_requirement",     "Điều kiện tốt nghiệp", "Graduation Requirements"),
    (_REF_PROG_HDR,  "reference_program",          "Chương trình tham khảo", "Reference Programs"),
    (_STAFF_HDR,     "staff_standard",             "Tiêu chuẩn giảng viên", "Staff Standards"),
    (_FACILITY_HDR,  "facility",                   "Cơ sở vật chất", "Facilities"),
    (_GUIDE_HDR,     "implementation_guideline",   "Hướng dẫn thực hiện", "Implementation Guidelines"),
]

_ADMISSION_HDR = re.compile(r"chuẩn\s*đầu\s*vào|đối\s*tượng\s*tuyển\s*sinh|admission\s*req", re.I)

def _parse_sections(doc: Document, objectives: list) -> list:
    """Thu thập các phần văn bản tường thuật làm ProgramSection."""
    sections = []
    order = 1

    # Objective section
    if objectives:
        general_obj = [o for o in objectives if o.objective_type == "general"]
        specific_obj = [o for o in objectives if o.objective_type == "specific"]
        if general_obj:
            sections.append(ProgramSection(
                section_code=f"OBJ_GENERAL",
                section_title_vi="Mục tiêu chung",
                section_title_en="General Objective",
                section_type="general_objective",
                content_vi=general_obj[0].description_vi,
                display_order=order,
            ))
            order += 1
        if specific_obj:
            combined = "\n".join(f"- {o.objective_code}: {o.description_vi}" for o in specific_obj)
            sections.append(ProgramSection(
                section_code="OBJ_SPECIFIC",
                section_title_vi="Mục tiêu cụ thể",
                section_title_en="Specific Objectives",
                section_type="specific_objective",
                content_vi=combined,
                display_order=order,
            ))
            order += 1

    # Scrape other narrative sections from paragraphs
    paras = list(doc.paragraphs)
    i = 0
    while i < len(paras):
        t = _c(paras[i].text)
        if not t:
            i += 1; continue

        for hdr_re, sec_type, title_vi, title_en in _SECTION_PATTERNS:
            if hdr_re.search(t) and len(t) < 100:
                # Collect following paragraphs until next major section
                chunks = []
                j = i + 1
                while j < len(paras):
                    pt = _c(paras[j].text)
                    if not pt:
                        j += 1; continue
                    if any(h.search(pt) for h, _, _, _ in _SECTION_PATTERNS) and len(pt) < 100:
                        break
                    if re.match(r"^\d+\.\s{1,3}[A-ZÁÀẢÃẠĂẮẶẰẴẶÊẾỆỀỂẼÔỐỘỒỔỖ]", pt) and len(pt) < 80:
                        break
                    chunks.append(pt)
                    j += 1

                if chunks:
                    sections.append(ProgramSection(
                        section_code=sec_type.upper()[:20],
                        section_title_vi=title_vi,
                        section_title_en=title_en,
                        section_type=sec_type,
                        content_vi=" ".join(chunks[:20]),
                        display_order=order,
                    ))
                    order += 1
                i = j
                break
        else:
            i += 1

    # Admission requirements
    collect = False
    adm_chunks = []
    for p in paras:
        t = _c(p.text)
        if not t: continue
        if _ADMISSION_HDR.search(t) and len(t) < 80:
            collect = True; continue
        if collect:
            if re.match(r"^\d+\.", t) and len(t) < 60: break
            if any(h.search(t) for h, _, _, _ in _SECTION_PATTERNS) and len(t) < 80: break
            if len(t) > 10: adm_chunks.append(t)
    if adm_chunks:
        sections.append(ProgramSection(
            section_code="ADMISSION_REQ",
            section_title_vi="Chuẩn đầu vào / Đối tượng tuyển sinh",
            section_title_en="Admission Requirements",
            section_type="admission_requirement",
            content_vi=" ".join(adm_chunks[:10]),
            display_order=order,
        ))

    return sections


# ── Components (program structure) ────────────────────────────────────────────

# Matches: "1.", "1.1.", "I.", "II.", "I.1." etc.
_COMP_CODE_RE = re.compile(r"^((?:[IVX]+|\d+)(?:\.(?:[IVX]+|\d+))*)\.\s+(.+)$")
# For normalizing Roman numeral codes to Arabic: I→1, II→2, III→3, IV→4...
_ROMAN_MAP = {"I":"1","II":"2","III":"3","IV":"4","V":"5","VI":"6","VII":"7","VIII":"8","IX":"9","X":"10"}
_SPECIAL_CREDIT_RE = re.compile(r"(\d+)\s*tiết|(\d+)\s*tc|(\d+)\s*tín\s*chỉ\s*quy\s*đổi", re.I)

def _parse_components(doc: Document) -> list:
    components = []

    for tbl in doc.tables:
        if not tbl.rows or len(tbl.rows) < 3: continue
        hdr_text = " ".join(_c(c.text).lower() for c in tbl.rows[0].cells)

        # Bảng cấu trúc: có "thành phần" hoặc "kiến thức" + "số tc"
        if not re.search(r"thành\s*phần|kiến\s*thức|số\s*tc|tín\s*chỉ", hdr_text, re.I):
            continue
        # Không phải bảng danh sách học phần
        # Bỏ qua bảng KHỐI LƯỢNG KIẾN THỨC TOÀN KHÓA (không có comp structure)
        if re.search(r"khối\s*lượng\s*kiến\s*thức\s*toàn", hdr_text, re.I): continue
        # Kiểm tra bảng phải có ít nhất 1 dòng bắt đầu bằng comp code (1., 1.1.)
        if not any(_COMP_CODE_RE.match(_c(row.cells[0].text)) for row in tbl.rows[1:8] if row.cells): continue
        if re.search(r"mã\s*hp|mã\s*học\s*phần", hdr_text, re.I): continue

        comp_col = -1
        credit_col = -1
        note_col = -1
        for i, h in enumerate(_c(c.text).lower() for c in tbl.rows[0].cells):
            if re.search(r"thành\s*phần|nội\s*dung|kiến\s*thức|component", h) and comp_col < 0:
                comp_col = i
            if re.search(r"số\s*tc|số\s*tín|credit|khối\s*lượng.*tín|^tín\s*chỉ", h) and credit_col < 0:
                credit_col = i
            if re.search(r"ghi\s*chú|note", h) and note_col < 0:
                note_col = i

        if comp_col < 0: comp_col = 0
        if credit_col < 0: credit_col = 1

        order = 1
        parent_stack = []  # Stack of (depth, code)

        for row in tbl.rows[1:]:
            cells = [_c(c.text) for c in row.cells]
            if not cells: continue
            name_raw = cells[comp_col] if comp_col < len(cells) else ""
            if not name_raw: continue

            credit_raw = cells[credit_col] if credit_col < len(cells) else ""
            note_raw = cells[note_col] if note_col >= 0 and note_col < len(cells) else ""

            # Parse component code + name
            m = _COMP_CODE_RE.match(name_raw)
            if not m: continue
            raw_code = m.group(1)
            # Normalize Roman numeral codes to Arabic: "I.1" → "1.1"
            parts = raw_code.split(".")
            norm_parts = [_ROMAN_MAP.get(p.upper(), p) for p in parts]
            code = ".".join(norm_parts)
            name = _c(m.group(2))

            vi, en = _split_vi_en(name)

            # Determine credit
            credit = _parse_credit(credit_raw)
            credit_text_raw = None
            counts_toward = True

            # Special non-credit items: GDTC, GDQP
            if re.search(r"gdtc|thể\s*chất|gdqp|quốc\s*phòng", name, re.I):
                counts_toward = False
                credit_text_raw = credit_raw

            # Credit text "165 tiết/8 tín chỉ quy đổi"
            if not credit and credit_raw and credit_raw not in ("-", "—"):
                sm = _SPECIAL_CREDIT_RE.search(credit_raw)
                if sm: credit_text_raw = credit_raw

            # Determine depth from code: "1" → depth 1, "1.1" → depth 2
            depth = len(code.split("."))

            # Find parent
            while parent_stack and parent_stack[-1][0] >= depth:
                parent_stack.pop()
            parent_code = parent_stack[-1][1] if parent_stack else None

            # Determine component_type
            if depth == 1:
                ctype = "program_group"
            elif depth == 2:
                ctype = "component_group"
            else:
                ctype = "component_subgroup"

            if re.search(r"tự\s*chọn|elective", name, re.I):
                ctype = "elective_group"
            elif re.search(r"bắt\s*buộc|mandatory|compulsory", name, re.I):
                ctype = "mandatory_group"
            elif re.search(r"khóa\s*luận|thesis|tốt\s*nghiệp", name, re.I):
                ctype = "thesis_group"
            elif re.search(r"gdtc|thể\s*chất|gdqp|quốc\s*phòng", name, re.I):
                ctype = "special_requirement"

            components.append(ProgramComponent(
                component_code=code,
                component_name_vi=vi,
                component_name_en=en,
                component_type=ctype,
                credits_required=credit,
                credit_text=credit_text_raw or (credit_raw if not credit else None),
                counts_toward_total_credits=counts_toward,
                note_text=note_raw or None,
                parent_code=parent_code,
                display_order=order,
            ))
            parent_stack.append((depth, code))
            order += 1

        if components: break

    return components


# ── Courses ───────────────────────────────────────────────────────────────────

_SEMESTER_HDR_RE = re.compile(r"^(?:học\s*kỳ|hk|semester)\s*(\d+)$|^(\d+)$", re.I)
_ROMAN_SEM_HDR   = re.compile(r"^(I{1,3}V?|V?I{0,3})$")

def _parse_courses(doc: Document, components: list) -> list:
    courses = []
    comp_codes = [c.component_code for c in components]

    for tbl in doc.tables:
        if not tbl.rows or len(tbl.rows) < 5: continue
        hdr_text = " ".join(_c(c.text).lower() for c in tbl.rows[0].cells)

        # Phải là bảng danh sách học phần
        if not re.search(r"mã\s*hp|mã\s*học\s*phần", hdr_text, re.I): continue
        if not re.search(r"số\s*tc|số\s*tín|tín\s*chỉ", hdr_text, re.I): continue

        # Parse header
        hdr_cells = [_c(c.text) for c in tbl.rows[0].cells]
        # Xử lý bảng có merged header row (2 rows header)
        sec_hdr = [_c(c.text) for c in tbl.rows[1].cells] if len(tbl.rows) > 1 else []

        # Tìm các cột cần thiết — 2-pass để tránh code_col bị overwrite bởi prereq
        stt_col = name_col = code_col = credit_col = prereq_col = lang_col = sem_col_start = -1
        semester_cols = {}  # col_idx -> semester_no

        # Pass 1: detect unambiguous cols (prereq, lang, sem, stt)
        for i, h in enumerate(hdr_cells):
            hl = h.lower()
            if re.search(r"^stt$|^tt$|^#$", hl): stt_col = i
            elif re.search(r"tiên\s*quyết|prerequisite", hl): prereq_col = i
            elif re.search(r"ngôn\s*ngữ|language", hl): lang_col = i
            elif re.search(r"học\s*kỳ|semester|bố\s*trí", hl): sem_col_start = i

        # Pass 2: name/code/credit — skip already-assigned cols
        for i, h in enumerate(hdr_cells):
            if i in (stt_col, prereq_col, lang_col, sem_col_start): continue
            hl = h.lower()
            # name_col: take LAST match — merged-header tables repeat col text;
            # the second (rightmost) col is the real-name column
            if re.search(r"nội\s*dung|học\s*phần|course\s*name|môn\s*học|ten\s*hp|tên\s*học|chương\s*trình\s*đào\s*tạo", hl):
                name_col = i
            elif re.search(r"mã\s*hp|mã\s*học\s*phần|course\s*code", hl) and code_col < 0:
                code_col = i   # FIRST match only — avoid prereq column overwrite
            elif re.search(r"số\s*tc|số\s*tín|^tc$|credits?$", hl) and credit_col < 0:
                credit_col = i
            elif _SEMESTER_HDR_RE.match(h) or _ROMAN_SEM_HDR.match(h):
                sm = _SEMESTER_HDR_RE.match(h)
                rm = _ROMAN_SEM_HDR.match(h)
                if sm:
                    sno = int(sm.group(1) or sm.group(2))
                elif rm:
                    ROMAN = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6, "VII": 7, "VIII": 8}
                    sno = ROMAN.get(h.upper(), 0)
                else:
                    sno = 0
                if 1 <= sno <= 12:
                    semester_cols[i] = sno

        # Fallback: check secondary header for semester numbers
        if not semester_cols and sec_hdr:
            for i, h in enumerate(sec_hdr):
                sm = _SEMESTER_HDR_RE.match(h)
                rm = _ROMAN_SEM_HDR.match(h)
                if sm:
                    sno = int(sm.group(1) or sm.group(2))
                    if 1 <= sno <= 12: semester_cols[i] = sno
                elif rm:
                    ROMAN = {"I":1,"II":2,"III":3,"IV":4,"V":5,"VI":6,"VII":7,"VIII":8}
                    sno = ROMAN.get(h.upper(), 0)
                    if sno: semester_cols[i] = sno

        # Detect name_col nếu chưa có (dạng "Nội dung và kế hoạch giảng dạy" ở cột 1)
        if name_col < 0:
            for i, h in enumerate(hdr_cells):
                if re.search(r"nội\s*dung|kế\s*hoạch|chương\s*trình", h, re.I) and i > 0:
                    name_col = i; break
            if name_col < 0: name_col = 1

        if code_col < 0 or credit_col < 0: continue

        # Detect start row (skip secondary header rows)
        data_start = 1
        for rr in range(1, min(3, len(tbl.rows))):
            sample = _c(tbl.rows[rr].cells[0].text) if tbl.rows[rr].cells else ""
            if re.search(r"^(stt|tt|tổng\s*số|total)", sample, re.I):
                data_start = rr + 1

        current_component = comp_codes[0] if comp_codes else "UNKNOWN"
        order = 1
        line_no_counter = 1

        for row in tbl.rows[data_start:]:
            cells = [_c(c.text) for c in row.cells]
            if not cells: continue

            stt_raw  = cells[stt_col] if stt_col >= 0 and stt_col < len(cells) else ""
            name_raw = cells[name_col] if name_col >= 0 and name_col < len(cells) else ""
            code_raw = cells[code_col] if code_col >= 0 and code_col < len(cells) else ""
            cr_raw   = cells[credit_col] if credit_col >= 0 and credit_col < len(cells) else ""

            # If name_raw is just a row-sequence number (≤3 digits), try the adjacent col
            if name_raw and re.match(r"^\d{1,3}$", name_raw.strip()):
                for nc in (name_col + 1, name_col - 1):
                    if 0 <= nc < len(cells) and nc not in (stt_col, code_col, credit_col, prereq_col, lang_col, sem_col_start) and len(cells[nc]) > 3:
                        name_raw = cells[nc]; break

            if not name_raw: continue

            # Detect section header rows (e.g. "1. GIÁO DỤC ĐẠI CƯƠNG", "I. KIẾN THỨC...")
            comp_m = _COMP_CODE_RE.match(name_raw)
            if comp_m and not code_raw:
                # This is a component header row → update current_component
                raw_code_m = comp_m.group(1)
                parts_m = raw_code_m.split(".")
                matched_code = ".".join(_ROMAN_MAP.get(p.upper(), p) for p in parts_m)
                if matched_code in comp_codes:
                    current_component = matched_code
                else:
                    # Try partial match by name prefix
                    header_name = _c(comp_m.group(2)).lower()
                    for cc in comp_codes:
                        comp_obj = next((c for c in components if c.component_code == cc), None)
                        if comp_obj and (header_name in comp_obj.component_name_vi.lower() or
                                         comp_obj.component_name_vi.lower() in header_name):
                            current_component = cc; break
                continue

            # Skip rows without a course code
            if not code_raw or code_raw in ("-", "—", "GDTC", "GDQP"):
                # Still might be GDTC/GDQP special rows
                if code_raw in ("GDTC", "GDQP") or re.search(r"gdtc|thể\s*chất|gdqp|quốc\s*phòng", name_raw, re.I):
                    is_phys = bool(re.search(r"gdtc|thể\s*chất", name_raw + code_raw, re.I))
                    is_def = bool(re.search(r"gdqp|quốc\s*phòng|an\s*ninh", name_raw + code_raw, re.I))
                    vi, en = _split_vi_en(name_raw)
                    cr = _parse_credit(cr_raw)
                    courses.append(ProgramCourse(
                        line_no=None,
                        component_code=current_component,
                        course_code_snapshot=code_raw or ("GDTC" if is_phys else "GDQP"),
                        course_name_vi_snapshot=vi,
                        course_name_en_snapshot=en,
                        credit_value=cr,
                        credit_text=cr_raw if not cr else None,
                        credit_unit_type="equivalent_credit",
                        is_required=True,
                        counts_toward_total_credits=False,
                        is_physical_education=is_phys,
                        is_defense_education=is_def,
                        display_order=order,
                    ))
                    order += 1
                continue

            # Skip summary/total rows (even if code_raw is set)
            if re.search(r"^tổng\s*số|^total\s*credit", name_raw, re.I):
                continue

            # Skip merged section-header rows:
            # Nhận biết: name_raw và code_raw là cùng 1 text (merged cell)
            # VÀ không trông giống mã học phần
            if code_raw and name_raw and (
                code_raw == name_raw
                or (len(code_raw) > 25 and not re.search(r"\b[A-Z]{2,}\d{2,}\b", code_raw))
            ):
                if _COMP_CODE_RE.match(name_raw) or not re.search(r"\d", code_raw):
                    # Cập nhật current_component nếu tên khớp component
                    for cc in comp_codes:
                        comp_obj = next((x for x in components if x.component_code == cc), None)
                        if comp_obj:
                            nm_clean = re.sub(r"\s*\([^)]*\)\s*$", "", name_raw).strip().lower()
                            if comp_obj.component_name_vi.lower() in nm_clean or nm_clean.startswith(comp_obj.component_name_vi.lower()[:10]):
                                current_component = cc; break
                    continue

            # Parse line number
            line_no = None
            if stt_raw and stt_raw.isdigit():
                line_no = int(stt_raw)
            else:
                line_no = line_no_counter
            line_no_counter += 1

            # Split VI/EN name
            vi, en = _split_vi_en(name_raw)

            # Credits
            credit = _parse_credit(cr_raw)
            credit_unit = "credit"
            counts_tc = True

            # Special courses
            is_thesis = bool(re.search(r"khóa\s*luận|luận\s*văn|thesis|tốt\s*nghiệp", name_raw, re.I))
            is_phys   = bool(re.search(r"thể\s*chất|gdtc", name_raw + code_raw, re.I))
            is_def    = bool(re.search(r"quốc\s*phòng|gdqp|an\s*ninh", name_raw + code_raw, re.I))
            if is_phys or is_def: counts_tc = False

            # Prerequisite
            prereq = None
            if prereq_col >= 0 and prereq_col < len(cells):
                pv = cells[prereq_col]
                if pv and pv not in ("-", "—", ""):
                    prereq = pv

            # Teaching language
            lang = None
            if lang_col >= 0 and lang_col < len(cells):
                lv = cells[lang_col]
                if lv: lang = lv

            # Semester allocations (from single "Học kỳ" text column)
            sem_allocs = []
            if sem_col_start >= 0 and sem_col_start < len(cells):
                sem_text = cells[sem_col_start]
                sem_nos = _roman_semester(sem_text)
                for sno in sem_nos:
                    sem_allocs.append(CourseSemesterAllocation(
                        semester_no=sno,
                        allocated_value=credit,
                        allocated_text=sem_text,
                    ))
            # Multi-column semester
            for col_i, sno in semester_cols.items():
                if col_i >= len(cells): continue
                val_raw = cells[col_i]
                if val_raw and val_raw not in ("-", "—", ""):
                    v = _parse_credit(val_raw)
                    sem_allocs.append(CourseSemesterAllocation(
                        semester_no=sno,
                        allocated_value=v,
                        allocated_text=val_raw if not v else None,
                    ))

            courses.append(ProgramCourse(
                line_no=line_no,
                component_code=current_component,
                course_code_snapshot=code_raw,
                course_name_vi_snapshot=vi,
                course_name_en_snapshot=en,
                credit_value=credit,
                credit_text=cr_raw if not credit else None,
                credit_unit_type=credit_unit,
                is_required=True,
                counts_toward_total_credits=counts_tc,
                is_thesis=is_thesis,
                is_physical_education=is_phys,
                is_defense_education=is_def,
                prerequisite_code=prereq,
                teaching_language=lang,
                semester_allocations=sem_allocs,
                display_order=order,
            ))
            order += 1

        if courses: break

    return courses


# ── Career paths ──────────────────────────────────────────────────────────────

def _parse_career_paths(doc: Document) -> list:
    paths = []
    collecting = False
    order = 1
    chunks_vi = []

    for p in doc.paragraphs:
        t = _c(p.text)
        if not t: continue

        if _CAREER_HDR.search(t):
            collecting = True
            # Inline: career info all in one long paragraph
            if len(t) > 80:
                chunks_vi.append(t)
            continue

        if collecting:
            # Stop at next major section
            if _GRAD_REQ_HDR.search(t) and len(t) < 80: break
            if _STRUCT_HDR.search(t) and len(t) < 80: break
            if re.match(r"^\d+\.\s{1,3}[A-ZÁÀẢÃẠ]", t) and len(t) < 80: break
            if len(t) > 20:
                chunks_vi.append(t)

    if chunks_vi:
        paths.append(CareerPath(
            description_vi=" ".join(chunks_vi[:20]),
            display_order=order,
        ))

    return paths


# ── Graduation requirements ───────────────────────────────────────────────────

_GRAD_REQ_KIND_MAP = [
    (re.compile(r"tín\s*chỉ|credits?", re.I),      "credits"),
    (re.compile(r"gpa|điểm\s*trung\s*bình", re.I), "gpa"),
    (re.compile(r"ngoại\s*ngữ|foreign\s*lang|tiếng\s*anh", re.I), "foreign_language"),
    (re.compile(r"tin\s*học|it\s*standard|công\s*nghệ\s*thông\s*tin", re.I), "it_standard"),
    (re.compile(r"thể\s*chất|gdtc|physical\s*edu", re.I), "physical_education"),
    (re.compile(r"quốc\s*phòng|gdqp|defense", re.I), "defense_education"),
    (re.compile(r"sinh\s*hoạt\s*công\s*dân|citizen", re.I), "citizen_week"),
    (re.compile(r"kỷ\s*luật|conduct", re.I),        "conduct_score"),
]

def _guess_grad_type(text: str) -> str:
    for pat, kind in _GRAD_REQ_KIND_MAP:
        if pat.search(text): return kind
    return "other"


def _parse_graduation_requirements(doc: Document) -> list:
    reqs = []
    collecting = False
    order = 1

    for p in doc.paragraphs:
        t = _c(p.text)
        if not t: continue
        if _GRAD_REQ_HDR.search(t) and len(t) < 80:
            collecting = True; continue
        if collecting:
            if re.match(r"^\d+\.\s{1,3}[A-ZÁÀẢÃẠ]", t) and not re.search(r"[àáâãèéêìíòóôõùú]", t[4:]) and len(t) < 80:
                break
            if _CAREER_HDR.search(t) and len(t) < 80: break
            if _STRUCT_HDR.search(t) and len(t) < 80: break
            cleaned = re.sub(r"^[-•*+\s]+", "", t).strip()
            if len(cleaned) > 15:
                reqs.append(GraduationRequirement(
                    requirement_type=_guess_grad_type(cleaned),
                    requirement_text_vi=cleaned,
                    display_order=order,
                ))
                order += 1

    return reqs


# ── Main entry ────────────────────────────────────────────────────────────────

def parse_curriculum(file_bytes: bytes, filename: str) -> CurriculumRecord:
    doc = Document(BytesIO(file_bytes))

    info       = _parse_general_info(doc)
    phil_vi, phil_en = _parse_philosophy(doc)
    objectives = _parse_objectives(doc)
    plos       = _parse_plos(doc)
    po_maps    = _parse_po_plo_map(doc, objectives, plos)
    sections   = _parse_sections(doc, objectives)
    components = _parse_components(doc)
    courses    = _parse_courses(doc, components)
    career_paths = _parse_career_paths(doc)
    grad_reqs  = _parse_graduation_requirements(doc)

    return CurriculumRecord(
        # Program info
        program_name_vi=info["program_name_vi"] or Path(filename).stem,
        program_name_en=info["program_name_en"],
        degree_name_vi=info["degree_name_vi"],
        degree_name_en=info["degree_name_en"],
        level_name_vi=info["level_name_vi"] or "Đại học",
        level_name_en=info["level_name_en"],

        # Version info
        version_label=info["version_label"] or "draft",
        major_code=info["major_code"],
        major_name_vi=info["major_name_vi"],
        major_name_en=info["major_name_en"],
        training_type_vi=info["training_type_vi"],
        training_type_en=info["training_type_en"],
        language_vi=info["language_vi"],
        language_en=info["language_en"],
        duration_text_vi=info["duration_text_vi"],
        duration_text_en=info["duration_text_en"],
        total_credits=info["total_credits"],
        applied_from_admission_year=info["applied_from_admission_year"],
        adjustment_time_text=info["adjustment_time_text"],
        program_opening_decision_text=info["program_opening_decision_text"],
        accreditation_text=info["accreditation_text"],
        degree_granting_unit_text=info["degree_granting_unit_text"],
        academic_management_unit_text=info["academic_management_unit_text"],
        philosophy_text_vi=phil_vi,
        philosophy_text_en=phil_en,
        issued_decision_text=info["issued_decision_text"],

        # Sub-tables
        objectives=objectives,
        plos=plos,
        po_plo_maps=po_maps,
        sections=sections,
        components=components,
        courses=courses,
        career_paths=career_paths,
        graduation_requirements=grad_reqs,
        reference_programs=[],

        source_file=filename,
    )
